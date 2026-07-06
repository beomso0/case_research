# -*- coding: utf-8 -*-
"""
판례 리서치 docx 빌더 (case-researcher Phase3).
- 최종본: 유용성 상위 N건 (판례번호 하이퍼링크 + 사실관계/설시 bold)
- 백업본: 전체 부합건
사용:  python scripts/build_docx.py <슬러그> [상위N=30]
전제:
  data/cases/<슬러그>/analysis/*.json  (부합 판정 rows; match/score/판례번호/유무죄/확정여부/내용_무죄부분/내용_유죄부분/비고)
  data/cases/<슬러그>/refined/*.json    (정밀 재평가 rows; fit/court/사실관계/설시 with **bold** 마커/...)  ← 최종본용(있으면)
  data/cases/<슬러그>/_manifest.json     (casenum→url; 하이퍼링크용)
  .claude/style_case_research.docx       (템플릿)
"""
import json, io, os, glob, re, sys
from docx import Document
from docx.oxml.shared import OxmlElement, qn

# Windows 콘솔 stdout이 cp949면 ™·→·✅·— 등 특수문자를 print할 때 UnicodeEncodeError(Exit 1)로 죽는다 → UTF-8 강제.
for _s in (sys.stdout, sys.stderr):
    try: _s.reconfigure(encoding="utf-8")
    except Exception: pass

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SLUG = sys.argv[1] if len(sys.argv) > 1 else "결과"
TOP_N = int(sys.argv[2]) if len(sys.argv) > 2 else 30
BASE = os.path.join(ROOT, "data", "cases", SLUG)
TEMPLATE = os.path.join(ROOT, ".claude", "style_case_research.docx")
OUTDIR = os.path.join(ROOT, "output"); os.makedirs(OUTDIR, exist_ok=True)
# 최종본(final docx) 저장 위치: 환경변수 RESEARCH_OUTPUT_DIR가 있으면 그곳, 없으면 repo output/.
# (사용자 규칙: 최종 결과물은 iCloud의 research_output에 모은다 — .claude/settings.local.json의 env로 주입.)
# 백업본·_final_ranking.json 등 작업 산출물은 항상 repo output/·data/에 남긴다.
FINAL_OUTDIR = (os.environ.get("RESEARCH_OUTPUT_DIR") or "").strip() or OUTDIR
try: os.makedirs(FINAL_OUTDIR, exist_ok=True)
except Exception: FINAL_OUTDIR = OUTDIR

# ---- url map ----
url_by_case = {}
try:
    for m in json.load(io.open(os.path.join(BASE, "_manifest.json"), encoding="utf-8")):
        if m.get("casenum") and m.get("url"):
            url_by_case[m["casenum"]] = m["url"]
except Exception:
    pass
date_re = re.compile(r'(\d{4})\.\s*\d{1,2}\.\s*\d{1,2}\.')
case_re = re.compile(r'(\d{2,4}[가-힣]{1,5}\d+(?:-\d+)?(?:-[A-Z])?)')
def make_url(pan):
    m = date_re.search(pan or "")
    if not m: return None
    court = pan[:m.start()].strip().replace(" ", "")
    m2 = re.search(r'(?:선고|자)\s*(.+?)\s*(?:판결|결정)', pan[m.end():])
    m3 = case_re.search(m2.group(1) if m2 else pan[m.end():])
    return "https://lbox.kr/case/%s/%s" % (court, m3.group(1)) if (court and m3) else None
def url_for(r):
    pan = r.get("판례번호") or ""
    m3 = case_re.search(pan)
    if m3 and m3.group(1) in url_by_case: return url_by_case[m3.group(1)]
    return make_url(pan)

def load(subdir):
    rows = []
    for fp in glob.glob(os.path.join(BASE, subdir, "*.json")):
        try: data = json.load(io.open(fp, encoding="utf-8"))
        except Exception: continue
        for r in (data.get("rows") if isinstance(data, dict) else data) or []:
            rows.append(r)
    return rows
def dedupe(rows, scorekey):
    best = {}
    for r in rows:
        k = re.sub(r"\.md$", "", str(r.get("file") or r.get("판례번호") or ""))
        if not k: continue
        if k not in best or (r.get(scorekey, 0) or 0) > (best[k].get(scorekey, 0) or 0): best[k] = r
    return list(best.values())

def court_bonus(c): return {"대법": 8, "고법": 4, "지법": 0}.get(c, 0)
def conf_bonus(c):
    c = c or ""; return 5 if ("확정" in c and "미확인" not in c) else 0

def add_hyperlink(p, url, text):
    rid = p.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hy = OxmlElement('w:hyperlink'); hy.set(qn('r:id'), rid)
    run = OxmlElement('w:r'); rPr = OxmlElement('w:rPr')
    for tag, val in (('w:color', '0563C1'), ('w:u', 'single')):
        e = OxmlElement(tag); e.set(qn('w:val'), val); rPr.append(e)
    run.append(rPr); t = OxmlElement('w:t'); t.text = text; run.append(t)
    hy.append(run); p._p.append(hy)
def add_rich(cell, text):
    for i, line in enumerate((text or "").split("\n")):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        if i == 0: p.text = ""
        for seg in re.split(r"(\*\*.+?\*\*)", line):
            if not seg: continue
            r = p.add_run(seg[2:-2] if seg.startswith("**") else seg); r.bold = seg.startswith("**")

def safe_save(doc, out):
    """대상 docx가 Word로 열려 있거나 iCloud 동기화로 잠겨(PermissionError) 있으면 _v2, _v3… 대체 경로로 저장."""
    try:
        doc.save(out); return out
    except PermissionError:
        base, ext = os.path.splitext(out)
        for i in range(2, 20):
            alt = "%s_v%d%s" % (base, i, ext)
            try:
                doc.save(alt)
                sys.stderr.write("[경고] %s 잠김(열려 있음/동기화 중) → %s 로 저장. Word 닫고 재실행하면 원본 갱신.\n"
                                 % (os.path.basename(out), os.path.basename(alt)))
                return alt
            except PermissionError:
                continue
        raise

def build(rows, sortkey, n, out, rich):
    doc = Document(TEMPLATE); t = doc.tables[0]
    for row in list(t.rows[1:]): row._element.getparent().remove(row._element)
    for r in sorted(rows, key=lambda r: -sortkey(r))[:n]:
        c = t.add_row().cells
        pan = r.get("판례번호") or r.get("file") or ""; url = url_for(r)
        p0 = c[0].paragraphs[0]; p0.text = ""
        (add_hyperlink(p0, url, pan) if url else p0.add_run(pan))
        c[1].text = r.get("유무죄") or ""; c[2].text = r.get("확정여부") or ""
        if rich:
            add_rich(c[3], "【사실관계】\n" + (r.get("사실관계") or "") + "\n\n【법원의 특징적 설시】\n" + (r.get("설시") or ""))
        else:
            mu = (r.get("내용_무죄부분") or "").strip(); yu = (r.get("내용_유죄부분") or "").strip()
            c[3].text = ("[무죄 부분] " + mu + ("\n\n[유죄 부분] " + yu if yu else "")) if mu else (r.get("reason") or "")
        c[4].text = r.get("비고") or ""
    saved = safe_save(doc, out); return len(t.rows) - 1, saved

# 최종본(정밀 재평가 refined/ 있으면 그걸로, 없으면 analysis/)
refined = dedupe(load("refined"), "fit")
if refined:
    comp = lambda r: (r.get("fit", 0) or 0) + court_bonus(r.get("court")) + conf_bonus(r.get("확정여부"))
    with io.open(os.path.join(BASE, "_final_ranking.json"), "w", encoding="utf-8") as f:
        json.dump([{"rank": i+1, "composite": comp(r), "url": url_for(r), **r} for i, r in enumerate(sorted(refined, key=lambda r: -comp(r)))], f, ensure_ascii=False, indent=1)
    nf, nf_path = build(refined, comp, TOP_N, os.path.join(FINAL_OUTDIR, "%s_판례리서치_최종.docx" % SLUG), rich=True)
    print("최종본:", nf, "행 (상위", TOP_N, ") →", nf_path)
# 백업본(전체 부합)
matched = [r for r in dedupe(load("analysis"), "score") if r.get("match")]
nb, nb_path = build(matched, lambda r: (r.get("score", 0) or 0), 10**6, os.path.join(OUTDIR, "%s_판례리서치.docx" % SLUG), rich=False)
print("백업본(전체 부합):", nb, "행 →", os.path.basename(nb_path))
